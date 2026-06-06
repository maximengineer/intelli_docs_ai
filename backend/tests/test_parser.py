import time
from io import BytesIO

import pytest
from app.core.errors import ParserTimeoutError
from app.documents.parser import parse_document, parse_document_with_timeout
from app.documents.schemas import ParsedDocument, ParsedPage
from docx import Document


def test_txt_parser_preserves_filename_and_text() -> None:
    parsed = parse_document(
        document_id="doc_test",
        filename="sample.txt",
        content=b"Invoice\nVendor: Acme Analytics Ltd\nTotal Amount: EUR 12,450.00",
    )

    assert parsed.document_id == "doc_test"
    assert parsed.filename == "sample.txt"
    assert "Acme Analytics" in parsed.text
    assert parsed.pages[0].page_number == 1


def test_docx_parser_extracts_paragraphs() -> None:
    buffer = BytesIO()
    document = Document()
    document.add_heading("Service Agreement", level=1)
    document.add_paragraph("Party: Northwind Retail Group")
    document.add_paragraph("Renewal Terms: Auto-renewal for one-year terms.")
    document.save(buffer)

    parsed = parse_document(
        document_id="doc_docx",
        filename="agreement.docx",
        content=buffer.getvalue(),
    )

    assert "Northwind Retail Group" in parsed.text
    assert "Auto-renewal" in parsed.text
    assert parsed.pages[0].page_number is None


def test_pdf_parser_extracts_digital_text() -> None:
    parsed = parse_document(
        document_id="doc_pdf",
        filename="invoice.pdf",
        content=_minimal_pdf("Invoice Total Amount EUR 18,300"),
    )

    assert "Invoice Total Amount EUR 18,300" in parsed.text
    assert parsed.pages[0].page_number == 1


def test_parser_timeout_raises_domain_error() -> None:
    def slow_parser(document_id: str, filename: str, content: bytes) -> ParsedDocument:
        del document_id, filename, content
        time.sleep(0.2)
        return ParsedDocument(
            document_id="doc_slow",
            filename="slow.txt",
            text="slow",
            pages=[ParsedPage(page_number=1, text="slow")],
        )

    with pytest.raises(ParserTimeoutError):
        parse_document_with_timeout(
            document_id="doc_slow",
            filename="slow.txt",
            content=b"slow",
            timeout_seconds=0.01,
            parser_func=slow_parser,
        )


def _minimal_pdf(text: str) -> bytes:
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode("ascii")
    objects = [
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n",
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n",
        (
            b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj\n"
        ),
        b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n",
        b"5 0 obj << /Length "
        + str(len(stream)).encode("ascii")
        + b" >> stream\n"
        + stream
        + b"\nendstream endobj\n",
    ]
    output = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for obj in objects:
        offsets.append(len(output))
        output.extend(obj)
    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode("ascii"))
    for offset in offsets:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    return bytes(output)
