from collections.abc import Callable
from concurrent.futures import ThreadPoolExecutor
from concurrent.futures import TimeoutError as FutureTimeoutError
from pathlib import Path
from tempfile import NamedTemporaryFile

from app.core.errors import ParserTimeoutError, UnsupportedFileTypeError
from app.documents.schemas import ParsedDocument, ParsedPage

ParserFunc = Callable[[str, str, bytes], ParsedDocument]


def parse_document_with_timeout(
    document_id: str,
    filename: str,
    content: bytes,
    timeout_seconds: float,
    parser_func: ParserFunc | None = None,
) -> ParsedDocument:
    parser = parser_func or parse_document
    executor = ThreadPoolExecutor(max_workers=1, thread_name_prefix="document-parser")
    future = executor.submit(parser, document_id, filename, content)
    try:
        return future.result(timeout=timeout_seconds)
    except FutureTimeoutError as exc:
        future.cancel()
        raise ParserTimeoutError(
            f"Parser timed out after {timeout_seconds:g} seconds for {filename}."
        ) from exc
    finally:
        executor.shutdown(wait=False, cancel_futures=True)


def parse_document(document_id: str, filename: str, content: bytes) -> ParsedDocument:
    suffix = Path(filename).suffix.lower()
    if suffix == ".txt":
        return _parse_txt(document_id, filename, content)
    if suffix == ".docx":
        return _parse_docx(document_id, filename, content)
    if suffix == ".pdf":
        return _parse_pdf(document_id, filename, content)
    raise UnsupportedFileTypeError(f"Unsupported file type: {suffix or 'unknown'}")


def _parse_txt(document_id: str, filename: str, content: bytes) -> ParsedDocument:
    text = content.decode("utf-8", errors="replace").strip()
    pages = [ParsedPage(page_number=1, text=text, section_title=_first_heading(text))]
    return ParsedDocument(document_id=document_id, filename=filename, text=text, pages=pages)


def _parse_docx(document_id: str, filename: str, content: bytes) -> ParsedDocument:
    try:
        from docx import Document
    except ImportError as exc:
        raise UnsupportedFileTypeError("DOCX parsing requires python-docx.") from exc

    with NamedTemporaryFile(suffix=".docx", delete=True) as tmp:
        tmp.write(content)
        tmp.flush()
        doc = Document(tmp.name)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    text = "\n".join(paragraphs).strip()
    pages = [ParsedPage(page_number=None, text=text, section_title=_first_heading(text))]
    return ParsedDocument(document_id=document_id, filename=filename, text=text, pages=pages)


def _parse_pdf(document_id: str, filename: str, content: bytes) -> ParsedDocument:
    try:
        import pdfplumber
    except ImportError as exc:
        raise UnsupportedFileTypeError("PDF parsing requires pdfplumber.") from exc

    pages: list[ParsedPage] = []
    with NamedTemporaryFile(suffix=".pdf", delete=True) as tmp:
        tmp.write(content)
        tmp.flush()
        with pdfplumber.open(tmp.name) as pdf:
            for index, page in enumerate(pdf.pages, start=1):
                page_text = (page.extract_text() or "").strip()
                if page_text:
                    pages.append(
                        ParsedPage(
                            page_number=index,
                            text=page_text,
                            section_title=_first_heading(page_text),
                        )
                    )

    text = "\n\n".join(page.text for page in pages).strip()
    if not text:
        raise UnsupportedFileTypeError("No digital text found. Scanned PDF OCR is out of scope.")
    return ParsedDocument(document_id=document_id, filename=filename, text=text, pages=pages)


def _first_heading(text: str) -> str | None:
    for line in text.splitlines():
        clean = line.strip()
        if clean and len(clean) <= 80:
            return clean
    return None
